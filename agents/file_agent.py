"""
FILE AGENT - FINAL PRO VERSION (MANIFEST-READY)
Generuje DOCX súbory na základe štruktúrovaného JSON manifestu.
Zabezpečuje, že ContentAgent a FileAgent hovoria rovnakým jazykom.
"""

import ast
import json
import logging
import re
import os
import sys
import zipfile
from datetime import datetime
from pathlib import Path

from config import PRODUCTS_DIR

log = logging.getLogger("FileAgent")

# Import-name -> pip package name (pre korektný requirements.txt)
_PKG_MAP = {
    "bs4": "beautifulsoup4", "PIL": "Pillow", "yaml": "PyYAML",
    "dotenv": "python-dotenv", "docx": "python-docx", "cv2": "opencv-python",
    "sklearn": "scikit-learn", "dateutil": "python-dateutil",
    "notion_client": "notion-client",
}

# DOCX imports
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    log.warning("⚠️ python-docx nie je nainštalovaný. Dokumenty budú v obmedzenom režime (TXT).")

# ══════════════════════════════════════════════════════════
#  JSON RESCUE ENGINE (Tvoj pôvodný, špičkový kód)
# ══════════════════════════════════════════════════════════════

def rescue_json(raw: str) -> dict:
    if not raw: return {}
    try:
        return json.loads(raw)
    except Exception: pass

    clean = re.sub(r"```(?:json)?\s*", "", raw).replace("```", "").strip()
    try:
        return json.loads(clean)
    except Exception: pass

    start = clean.find("{")
    end = clean.rfind("}") + 1
    if start >= 0 and end > start:
        candidate = clean[start:end]
        try:
            return json.loads(candidate)
        except Exception: pass
        try:
            return json.loads(candidate, strict=False)
        except Exception: pass

    if start >= 0 and end > start:
        fixed = _fix_common_json_errors(candidate)
        try:
            return json.loads(fixed, strict=False)
        except Exception: pass

    log.warning("⚠️ JSON rescue zlyhal, používam fallback.")
    return _heuristic_fallback(raw)

def _fix_common_json_errors(text: str) -> str:
    text = re.sub(r",\s*([}\]])", r"\1", text)
    text = re.sub(r':\s*"([^"]*)\n', r': "\1",\n', text)
    text = re.sub(r'(\{|,)\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*:', r'\1 "\2":', text)
    return text

def _heuristic_fallback(raw: str) -> dict:
    result = {"raw_content": raw}
    title_match = re.search(r'"title"\s*:\s*"([^"]+)"', raw)
    if not title_match:
        title_match = re.search(r'title[:\s]+([^\n]{5,60})', raw, re.IGNORECASE)
    if title_match:
        result["title"] = title_match.group(1).strip()
    return result

# ══════════════════════════════════════════════════════════════════
#  FILE AGENT
# ══════════════════════════════════════════════════════════════════════════════

class FileAgent:
    def save(self, product: dict) -> dict[str, Path]:
        """Hlavný vstupný bod: Uloží JSON, vygeneruje DOCX (DIY) a ZIP (PRO)."""

        # 1. JSON Rescue
        if "raw_content" in product and isinstance(product["raw_content"], str):
            rescued = rescue_json(product["raw_content"])
            if rescued:
                if "content_structure" in rescued or "pro_kit" in rescued:
                    product.update(rescued)
                    product.pop("raw_content", None)
                    log.info("✅ raw_content úspešne zachránený (plná štruktúra)")
                else:
                    product.update(rescued)
                    log.warning("⚠️ raw_content zachránený len čiastočne (použije sa fallback text)")

        # Extrahuj title z manifest štruktúry (metadata.title > title > fallback)
        meta = product.get("metadata", {})
        title = meta.get("title") or product.get("title") or "Digital Product"
        # Oprav slug-style title ("my-product-name" -> "My Product Name")
        if "-" in title and " " not in title:
            title = title.replace("-", " ").title()
        product["title"] = title  # Normalizuj pre downstream
        safe = self._safe_filename(title)

        # Vytvorenie štruktúry priečinkov
        root_folder = PRODUCTS_DIR / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{safe}"
        root_folder.mkdir(parents=True, exist_ok=True)

        diy_folder = root_folder / "v1_DIY_Basic"
        diy_folder.mkdir(exist_ok=True)

        pro_folder = root_folder / "v2_PRO_Automation_Kit"
        pro_folder.mkdir(exist_ok=True)

        # Uložíme čisté dáta
        (root_folder / "data.json").write_text(json.dumps(product, ensure_ascii=False, indent=2), encoding="utf-8")

        # 2. GENERÁCIA DOCX (Pre oboch verziami)
        docx_name = f"{safe}_Guide.docx"
        if DOCX_AVAILABLE:
            try:
                # DIY verzia (len dokument)
                self._generate_docx(product, diy_folder / docx_name)
                # PRO verzia (dokument + kód)
                self._generate_docx(product, pro_folder / docx_name)
                log.info(f"✅ DOCX vygenerované pre obe verzie: {docx_name}")
            except Exception as e:
                log.error(f"❌ DOCX generovanie zlyhalo: {e}")
                self._generate_basic_txt(product, diy_folder / f"{safe}.txt")
        else:
            self._generate_basic_txt(product, diy_folder / f"{safe}.txt")

        # 3. PRO EXTRA SÚBORY (ZIP BALÍČEK)
        if "pro_kit" in product:
            try:
                self._create_pro_zip(product, pro_folder, docx_name)
                log.info("✅ PRO ZIP balíček vytvorený.")
            except Exception as e:
                log.error(f"❌ Zlyhalo vytváraním PRO ZIPu: {e}")

        return {"json": root_folder / "data.json"}

    def _create_pro_zip(self, product, pro_folder, docx_name):
        """Zabalí DOCX + VŠETKY skripty + requirements + README + setup do plug-and-play ZIPu."""
        zip_path = pro_folder / f"{docx_name.replace('.docx', '')}_PRO_Kit.zip"

        # 1. Pozbieraj všetky kódové bloky (pro_kit + content_structure)
        blocks = self._collect_code_blocks(product)
        # 2. Zlaď lokálne importy so súbormi (zamedzí ModuleNotFoundError)
        self._reconcile_local_imports(blocks)
        # 3. Odvoď requirements zo skutočných importov + deklarovaných
        reqs = self._detect_requirements(product, blocks)
        # 4. Nájdi CLI vstupný skript pre README
        entry = self._find_entry_script(blocks)

        # 5. Verifikácia: žiadny lokálny import nesmie zostať bez súboru
        unresolved = self._verify_imports(blocks)
        if unresolved:
            log.warning(f"⚠️ Nevyriešené lokálne importy v ZIP balíčku: {unresolved}")

        # 6. Vygeneruj dokumentáciu a setup skripty
        readme = self._generate_readme(product, blocks, reqs, entry)
        setup_sh = self._generate_setup_sh()
        setup_bat = self._generate_setup_bat()

        # 7. Zápis ZIPu priamo cez writestr (bez dočasných súborov)
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.write(pro_folder / docx_name, arcname=docx_name)
            for b in blocks:
                zipf.writestr(b["name"], self._strip_code_fences(b["content"]))
            if reqs:
                zipf.writestr("requirements.txt", "\n".join(reqs) + "\n")
            zipf.writestr("README.md", readme)
            zipf.writestr("setup.sh", setup_sh)
            zipf.writestr("setup.bat", setup_bat)

        extra = ", requirements.txt" if reqs else ""
        log.info(f"📦 ZIP: {[b['name'] for b in blocks]} + README.md, setup.sh, setup.bat{extra}")

    # ──────────────────────────────────────────────────────────
    #  PRO ZIP HELPERY (packaging, AST, dokumentácia)
    # ──────────────────────────────────────────────────────────

    def _parse_ast(self, content: str):
        """Bezpečne sparsuje Python kód na AST (None pri syntax error)."""
        try:
            return ast.parse(content)
        except Exception:
            return None

    def _stdlib(self) -> set:
        return {m.lower() for m in getattr(sys, "stdlib_module_names", set())}

    def _module_symbols(self, content: str) -> set:
        """Top-level symboly (funkcie, triedy, premenné), ktoré modul definuje."""
        tree = self._parse_ast(content)
        syms = set()
        if not tree:
            return syms
        for node in tree.body:
            if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
                syms.add(node.name)
            elif isinstance(node, ast.Assign):
                for t in node.targets:
                    if isinstance(t, ast.Name):
                        syms.add(t.id)
        return syms

    def _import_from_targets(self, content: str) -> list:
        """Vráti [(module, [names])] pre 'from X import ...' (vrátane relatívnych)."""
        tree = self._parse_ast(content)
        out = []
        if not tree:
            return out
        for node in ast.walk(tree):
            if isinstance(node, ast.ImportFrom) and node.module:
                out.append((node.module, [a.name for a in node.names]))
        return out

    def _collect_code_blocks(self, product: dict) -> list:
        """Pozbiera VŠETKY .py bloky: pro_kit.scripts + content_structure[type=code]."""
        blocks = []
        seen = set()
        used_names = set()

        def _add(name: str, content: str):
            content = self._strip_code_fences(content)
            if not content.strip():
                return
            key = re.sub(r"\s+", "", content)
            if key in seen:
                return
            if not name.endswith(".py"):
                name = name + ".py"
            base = name[:-3]
            i = 2
            while name in used_names:
                name = f"{base}_{i}.py"
                i += 1
            seen.add(key)
            used_names.add(name)
            blocks.append({"name": name, "content": content})

        pro_kit = product.get("pro_kit", {})
        for s in pro_kit.get("scripts", []):
            name = s.get("name") or self._safe_filename(s.get("description", "script"))
            _add(name, s.get("content", ""))

        for item in product.get("content_structure", []):
            if item.get("type") != "code":
                continue
            lang = (item.get("language") or "python").lower()
            if lang not in ("python", "py"):
                continue
            _add(self._safe_filename(item.get("title", "module")) or "module", item.get("content", ""))

        return blocks

    def _reconcile_local_imports(self, blocks: list) -> None:
        """Premenuje súbory tak, aby 'from X import sym' sedelo na súbor, kde je sym definovaný."""
        sym_to_block = {}
        for b in blocks:
            for s in self._module_symbols(b["content"]):
                sym_to_block.setdefault(s, b)
        basenames = {b["name"][:-3] for b in blocks}

        for b in blocks:
            for module, names in self._import_from_targets(b["content"]):
                top = module.split(".")[0]
                if not top or top in basenames:
                    continue
                target = None
                for n in names:
                    cand = sym_to_block.get(n)
                    if cand and cand is not b:
                        target = cand
                        break
                if target:
                    old = target["name"]
                    new = f"{top}.py"
                    if old != new and new not in basenames:
                        log.info(f"🔗 Zlaďujem lokálny import: '{old}' → '{new}'")
                        basenames.discard(old[:-3])
                        target["name"] = new
                        basenames.add(top)

    def _verify_imports(self, blocks: list) -> list:
        """Nájde lokálne importy bez zodpovedajúceho súboru (genuine ModuleNotFoundError riziko)."""
        sym_to_block = {}
        for b in blocks:
            for s in self._module_symbols(b["content"]):
                sym_to_block.setdefault(s, b)
        basenames = {b["name"][:-3] for b in blocks}
        unresolved = []
        for b in blocks:
            for module, names in self._import_from_targets(b["content"]):
                top = module.split(".")[0]
                if top in basenames:
                    continue
                if any(n in sym_to_block for n in names):
                    unresolved.append(f"{b['name']}: from {module} import {', '.join(names)}")
        return unresolved

    def _detect_requirements(self, product: dict, blocks: list) -> list:
        """Odvodí requirements zo skutočných importov + deklarovaných (filtruje stdlib a lokálne)."""
        stdlib = self._stdlib()
        basenames = {b["name"][:-3] for b in blocks}
        found = set()
        for b in blocks:
            tree = self._parse_ast(b["content"])
            if not tree:
                continue
            for node in ast.walk(tree):
                if isinstance(node, ast.Import):
                    for a in node.names:
                        found.add(a.name.split(".")[0])
                elif isinstance(node, ast.ImportFrom) and node.module and node.level == 0:
                    found.add(node.module.split(".")[0])

        pkgs = []
        for mod in sorted(found):
            if not mod or mod.lower() in stdlib or mod in basenames:
                continue
            pkgs.append(_PKG_MAP.get(mod, mod))

        declared = list(product.get("pro_kit", {}).get("requirements", []))
        return self._filter_requirements(declared + pkgs)

    def _find_entry_script(self, blocks: list):
        """Nájde CLI vstupný skript (preferuje argparse + __main__)."""
        best, best_score = None, -1
        for b in blocks:
            c = b["content"]
            score = 0
            if "argparse" in c or "ArgumentParser" in c:
                score += 2
            if "__main__" in c:
                score += 1
            if "def main" in c:
                score += 1
            if score > best_score:
                best, best_score = b, score
        return best if best_score > 0 else None

    def _extract_cli(self, content: str):
        """Best-effort extrakcia subkomand a argumentov z argparse kódu."""
        subs = re.findall(r"add_parser\(\s*['\"]([^'\"]+)['\"]", content)
        args = re.findall(r"add_argument\(\s*['\"](--?[^'\"]+)['\"]", content)
        seen = set()
        args = [a for a in args if not (a in seen or seen.add(a))]
        return subs, args

    def _generate_readme(self, product: dict, blocks: list, reqs: list, entry) -> str:
        meta = product.get("metadata", {})
        title = meta.get("title") or product.get("title") or "Premium Python Kit"
        desc = product.get("marketing", {}).get("description") or meta.get("description") or ""
        entry_name = entry["name"] if entry else (blocks[0]["name"] if blocks else "main.py")

        # Usage príklady
        usage_lines = []
        if entry:
            subs, args = self._extract_cli(entry["content"])
            if subs:
                for s in subs:
                    usage_lines.append(f"python {entry_name} {s}")
            elif args:
                usage_lines.append(f"python {entry_name} " + " ".join(f"{a} <value>" for a in args[:4]))
            else:
                usage_lines.append(f"python {entry_name} --help")
        usage = "\n".join(usage_lines) if usage_lines else f"python {entry_name} --help"

        files_list = "\n".join(f"- `{b['name']}`" for b in blocks)
        req_note = "All dependencies are installed automatically by the setup script." if reqs \
            else "No external dependencies required (pure standard library)."

        return (
            f"# {title}\n\n"
            f"{desc}\n\n"
            "## 📦 What's Inside\n\n"
            f"{files_list}\n"
            "- `requirements.txt` — Python dependencies\n"
            "- `setup.sh` / `setup.bat` — one-click environment setup\n\n"
            "## 🚀 Quick Start / Installation\n\n"
            f"{req_note}\n\n"
            "**Windows:**\n\n"
            "```bat\n"
            "setup.bat\n"
            "```\n\n"
            "**macOS / Linux:**\n\n"
            "```bash\n"
            "chmod +x setup.sh\n"
            "./setup.sh\n"
            "```\n\n"
            "This creates a virtual environment (`.venv`) and installs all dependencies.\n\n"
            "## 🛠️ Usage Guide\n\n"
            "Activate the environment, then run the tool:\n\n"
            "**Windows:** `.venv\\Scripts\\activate`  |  **macOS/Linux:** `source .venv/bin/activate`\n\n"
            "```bash\n"
            f"{usage}\n"
            "```\n\n"
            "## 💡 Pro Tips\n\n"
            "- **Use proxies** when scraping at scale to avoid IP bans and rate limits. "
            "Rotate residential proxies for best results.\n"
            "- Run small test batches first to validate your target sources.\n\n"
            "## ⚠️ Disclaimer\n\n"
            "This software is provided **AS IS**, without warranty of any kind. "
            "You are solely responsible for complying with **GDPR** and all applicable data-protection "
            "and anti-spam laws in your jurisdiction. Target website structures (HTML) may change over "
            "time, which can require updates to the scraping logic.\n"
        )

    def _generate_setup_sh(self) -> str:
        return (
            "#!/usr/bin/env bash\n"
            "set -e\n"
            'echo "Creating virtual environment (.venv)..."\n'
            'if command -v python3 >/dev/null 2>&1; then PY=python3; else PY=python; fi\n'
            'if ! command -v "$PY" >/dev/null 2>&1; then\n'
            '  echo "[ERROR] Python was not found. Install it from https://www.python.org/downloads/"\n'
            "  exit 1\n"
            "fi\n"
            '"$PY" -m venv .venv\n'
            'echo "Upgrading pip..."\n'
            ".venv/bin/python -m pip install --upgrade pip\n"
            "if [ -f requirements.txt ]; then\n"
            '  echo "Installing dependencies..."\n'
            "  .venv/bin/python -m pip install -r requirements.txt\n"
            "fi\n"
            'echo "[OK] Setup complete! Activate anytime with: source .venv/bin/activate"\n'
        )

    def _generate_setup_bat(self) -> str:
        return (
            "@echo off\r\n"
            "setlocal\r\n"
            "echo ============================================\r\n"
            "echo   Premium Kit - Automatic Setup\r\n"
            "echo ============================================\r\n"
            "\r\n"
            "REM --- Detect Python (prefer the 'py' launcher) ---\r\n"
            'set "PYEXE="\r\n'
            'py -3 --version >nul 2>&1 && set "PYEXE=py -3"\r\n'
            "if not defined PYEXE (\r\n"
            '  python --version >nul 2>&1 && set "PYEXE=python"\r\n'
            ")\r\n"
            "if not defined PYEXE (\r\n"
            "  echo [ERROR] Python was not found.\r\n"
            "  echo Install it from https://www.python.org/downloads/ and tick \"Add Python to PATH\".\r\n"
            "  pause\r\n"
            "  exit /b 1\r\n"
            ")\r\n"
            "\r\n"
            "echo Using Python: %PYEXE%\r\n"
            "echo Creating virtual environment (.venv)...\r\n"
            "%PYEXE% -m venv .venv\r\n"
            "if errorlevel 1 (\r\n"
            "  echo [ERROR] Failed to create the virtual environment.\r\n"
            "  pause\r\n"
            "  exit /b 1\r\n"
            ")\r\n"
            "\r\n"
            "echo Upgrading pip...\r\n"
            '".venv\\Scripts\\python.exe" -m pip install --upgrade pip\r\n'
            "\r\n"
            "if exist requirements.txt (\r\n"
            "  echo Installing dependencies...\r\n"
            '  ".venv\\Scripts\\python.exe" -m pip install -r requirements.txt\r\n'
            "  if errorlevel 1 (\r\n"
            "    echo [ERROR] Failed to install dependencies.\r\n"
            "    pause\r\n"
            "    exit /b 1\r\n"
            "  )\r\n"
            ")\r\n"
            "\r\n"
            "echo.\r\n"
            "echo [OK] Setup complete!\r\n"
            "echo To activate later, run: .venv\\Scripts\\activate.bat\r\n"
            "pause\r\n"
        )

    def _generate_docx(self, product, output_path):
        """Generuje DOCX na základe nového MANIFEST štruktúry."""
        document = Document()

        # --- ŠTÝLY ---
        styles = document.styles
        self._setup_styles(styles)

        # 1. Titulok (z metadata)
        meta = product.get("metadata", {})
        title_text = meta.get("title") or product.get("title") or "Digital Product"
        document.add_paragraph(title_text, style='DocTitle')

        # 2. Popis (z marketing)
        marketing = product.get("marketing", {})
        desc = marketing.get("description") or product.get("description")
        if desc:
            document.add_paragraph(desc, style='BodyTextCustom')

        # 3. HLAVNÝ OBSAH (Iterácia cez content_structure)
        content_items = product.get("content_structure", [])
        for item in content_items:
            itype = item.get("type")

            if itype == 'heading':
                level = item.get("level", 1)
                document.add_heading(item.get("text", ""), level=level)

            elif itype == 'text':
                document.add_paragraph(item.get("text", ""), style='BodyTextCustom')

            elif itype == 'prompt':
                p = document.add_paragraph()
                run = p.add_run(f"💡 PROMPT: {item.get('title', 'Strategic Prompt')}")
                run.bold = True
                document.add_paragraph(item.get("content", ""), style='CodeStyle')

            elif itype == 'code':
                document.add_paragraph(item.get("title", "Code Snippet"), style='SectionH1')
                code_para = document.add_paragraph()
                code_content = self._strip_code_fences(item.get("content", ""))
                run = code_para.add_run(code_content)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                # Pridanie šedej pozadiny pre kód (ako si mal)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:val'), 'clear')
                shading_elm.set(qn('w:fill'), 'F7FAFC')
                code_para.paragraph_format.element.get_or_add_pPr().append(shading_elm)

        # 4. Fallback (ak je content_structure prázdna)
        if not content_items and product.get("raw_content"):
            document.add_heading("Product Details", level=1)
            document.add_paragraph(product["raw_content"], style='BodyTextCustom')

        document.save(output_path)

    def _setup_styles(self, styles):
        """Konfigurácia vizuálneho štýlu dokumentu."""
        # Title
        title_style = styles.add_style('DocTitle', 1)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(24)
        title_style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # H1
        h1_style = styles.add_style('SectionH1', 1)
        h1_style.font.name = 'Arial'
        h1_style.font.size = Pt(16)
        h1_style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        h1_style.paragraph_format.space_before = Pt(12)

        # Body
        body_style = styles.add_style('BodyTextCustom', 1)
        body_style.font.name = 'Arial'
        body_style.font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(6)

        # Code
        code_style = styles.add_style('CodeStyle', 1)
        code_style.font.name = 'Courier New'
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Inches(0.2)

    def _generate_basic_txt(self, product, output_path):
        content = f"{product.get('title', 'Product')}\n\n{product.get('description', '')}\n\n{product.get('raw_content', '')}"
        output_path.write_text(content, encoding="utf-8")

    def _filter_requirements(self, reqs: list) -> list:
        """Odstráni stdlib moduly z requirements (nedajú sa pip-installnúť) a duplikáty."""
        import sys
        stdlib = {m.lower() for m in getattr(sys, "stdlib_module_names", set())}
        seen = set()
        out = []
        for r in reqs:
            if not r or not isinstance(r, str):
                continue
            name = re.split(r"[<>=!\[ ]", r.strip(), maxsplit=1)[0].strip().lower()
            if not name or name in stdlib or name in seen:
                continue
            seen.add(name)
            out.append(r.strip())
        return out

    def _strip_code_fences(self, text: str) -> str:
        """Odstráni markdown code fences (```python, ```py, ```) z textu."""
        if not text:
            return text
        text = re.sub(r'```(?:python|py)?\s*\n?', '', text)
        text = text.replace('```', '')
        return text.strip()

    def _safe_filename(self, text: str, max_len: int = 40) -> str:
        safe = re.sub(r'[^\w\s-]', '', text.lower())
        safe = re.sub(r'[\s-]+', '_', safe)
        return safe[:max_len].strip("_")